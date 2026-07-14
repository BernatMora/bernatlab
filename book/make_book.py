# -*- coding: utf-8 -*-
"""
make_book.py — Generador de BernatLab_Manual_Modul_1 (PDF + DOCX) a partir dels capítols Markdown.

Estratègia: fem un parser Markdown propi, simple i robust, orientat a les
construccions que usem al manual (capçaleres, paràgrafs, llistes, codi,
cites, taules, regles horitzontals, bloc Mermaid). Sortida: PDF amb reportlab
i DOCX amb python-docx.

Aquest script NO és un conversor Markdown universal. És un conversor
específic per a l'estil d'aquest llibre. Si més endavant afegim sintaxi
nova, l'ampliem aquí.
"""

import os
import re
import sys
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Paragraph,
    Preformatted,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    KeepTogether,
    NextPageTemplate,
)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Configuració
# ---------------------------------------------------------------------------

ROOT = Path(__file__).resolve().parent
# Capítols són a book/chapters/ — el script pot córrer des de qualsevol lloc
# gràcies a aquesta resolució relativa al __file__.
CAP_DIR = ROOT / "chapters"

CHAPTERS_M1 = [
    ("01-que-es-bernatlab.md", "Capítol 1 — Què és BernatLab"),
    ("02-raspberry-pi.md", "Capítol 2 — La Raspberry Pi 4 per dins"),
    ("03-linux-servidor.md", "Capítol 3 — Linux per administrar un servidor"),
    ("04-xarxa-ssh-tailscale.md", "Capítol 4 — Xarxa, SSH i Tailscale"),
    ("05-docker.md", "Capítol 5 — Docker des de zero"),
    ("06-portainer.md", "Capítol 6 — Portainer"),
    ("07-uptime-kuma.md", "Capítol 7 — Uptime Kuma"),
    ("08-homepage.md", "Capítol 8 — Homepage"),
    ("09-git-documentacio.md", "Capítol 9 — Git i documentació"),
    ("10-full-de-ruta.md", "Capítol 10 — Full de ruta del BernatLab"),
]

CHAPTERS_M2 = [
    ("11-del-m1-al-m2.md", "Capítol 11 — Del Mòdul 1 al M2: què construïm"),
    ("12-mqtt-des-de-zero.md", "Capítol 12 — MQTT des de zero"),
    ("13-mosquitto-bernatlab.md", "Capítol 13 — Mosquitto al BernatLab"),
    ("14-publicar-dades-sensors.md", "Capítol 14 — Publicar dades: els sensors"),
    ("15-influxdb.md", "Capítol 15 — InfluxDB: base de dades de sèries temporals"),
    ("16-telegraf.md", "Capítol 16 — Telegraf: el pont"),
    ("17-node-red.md", "Capítol 17 — Node-RED: programació visual"),
    ("18-fluxos-practics.md", "Capítol 18 — Fluxos pràctics"),
    ("19-grafana.md", "Capítol 19 — Grafana: visualitzar les dades"),
    ("20-api-publica.md", "Capítol 20 — API pública: servir les dades al món"),
    ("21-integracio-hort-osona.md", "Capítol 21 — Integració amb Hort Osona"),
    ("22-operativa.md", "Capítol 22 — Operativa: còpies, alertes, escalat"),
]

CHAPTERS_M3 = [
    ("23-que-es-lora.md", "Capítol 23 — Què és LoRa i per què per a Hort Osona"),
    ("24-fisica-radio.md", "Capítol 24 — Física de ràdio: els paràmetres que importen"),
    ("25-lorawan-vs-p2p.md", "Capitol 25 — LoRaWAN vs LoRa P2P: l'arbre de decisió"),
    ("26-lorawan-ttn.md", "Capitol 26 — La capa LoRaWAN: TTN, device profiles, payloads"),
    ("27-gateway-raspberry.md", "Capitol 27 — Gateway LoRaWAN a la Raspberry amb Concentratord"),
    ("28-node-hardware.md", "Capitol 28 — El node: ESP32 + SX1262, hardware i esquemes"),
    ("29-programacio-node.md", "Capitol 29 — Programació del node: ESP32 + LoRaWAN"),
    ("30-recepcio-bernatlab.md", "Capitol 30 — Recepció al BernatLab: de TTN a InfluxDB"),
    ("31-lora-p2p.md", "Capitol 31 — LoRa P2P amb SX1262: xarxes privades sense TTN"),
    ("32-proves-camp.md", "Capitol 32 — Proves de camp, cobertura i resolució de problemes"),
]

CHAPTERS_M4 = [
    ("33-que-es-ia-local.md", "Capitol 33 — Què és la IA local i per què a Hort Osona"),
    ("34-ollama-instalacio.md", "Capitol 34 — Ollama al Mac: instal·lació i primera conversa"),
    ("35-escollir-model.md", "Capitol 35 — Com triar el millor model: mida, velocitat, qualitat, català"),
    ("36-embeddings-vectorstore.md", "Capitol 36 — Embeddings i bases vectorials"),
    ("37-rag-pas-a-pas.md", "Capitol 37 — RAG pas a pas: carregar les 76 fitxes d'hort a Ollama"),
    ("38-client-web.md", "Capitol 38 — Client web"),
    ("39-api-ollama.md", "Capitol 39 — API d'Ollama: integrar la IA amb el BernatLab"),
    ("40-veu.md", "Capitol 40 — Veu: parlar a l'assistent"),
    ("41-privadesa.md", "Capitol 41 — Privadesa i bones pràctiques"),
    ("42-casos-us.md", "Capitol 42 — 10 consultes reals a l'assistent Hort Osona"),
]

CHAPTERS_M5 = [
    ("43-filosofia-seguretat.md", "Capitol 43 — Filosofia de seguretat al BernatLab"),
    ("44-tailscale-acls.md", "Capitol 44 — Tailscale ACLs i segmentació de xarxa"),
    ("45-copies-seguretat.md", "Capitol 45 — Còpies de seguretat amb restic i BorgBackup"),
    ("46-2fa-secrets.md", "Capitol 46 — 2FA, secrets i gestió de claus"),
    ("47-fail2ban-tallafocs.md", "Capitol 47 — fail2ban, rate limiting i tallafocs aplicat"),
    ("48-hardening-so.md", "Capitol 48 — Hardening del sistema operatiu"),
    ("49-auditoria-logs.md", "Capitol 49 — Auditoria, logs de seguretat i resposta a incidents"),
    ("50-drp-recuperacio.md", "Capitol 50 — DRP: pla de recuperació davant desastres"),
]

CHAPTERS_M6 = [
    ("51-filosofia-operativa.md", "Capitol 51 — Filosofia operativa: del DIY al servei 24/7"),
    ("52-monitoratge-grafana.md", "Capitol 52 — Monitoratge avançat amb Grafana i Prometheus"),
    ("53-alertes-telegram.md", "Capitol 53 — Alertes intel·ligents amb Grafana i Telegram"),
    ("54-scripts-manteniment.md", "Capitol 54 — Scripts de manteniment i actualitzacions"),
    ("55-runbooks.md", "Capitol 55 — Runbooks: procediments pas a pas"),
    ("56-diagnostic-troubleshooting.md", "Capitol 56 — Diagnòstic i troubleshooting 24/7"),
    ("57-pujar-hardware.md", "Capitol 57 — Quan cal pujar de hardware: escenaris reals"),
]

CHAPTERS_M7 = [
    ("58-preparacio.md", "Capitol 58 — Preparació: què necessites tenir abans de començar"),
    ("59-primer-contacte.md", "Capitol 59 — Primer contacte amb la Raspberry Pi"),
    ("60-sistema-base-segur.md", "Capitol 60 — Sistema base segur"),
    ("61-docker-portainer.md", "Capitol 61 — Docker i Portainer: la base dels serveis"),
    ("62-uptime-kuma.md", "Capitol 62 — Uptime Kuma: el primer monitor"),
    ("63-mqtt-influxdb-grafana.md", "Capitol 63 — La cadena de dades: MQTT, InfluxDB, Grafana"),
    ("64-node-red.md", "Capitol 64 — Node-RED: les primeres automatitzacions"),
    ("65-node-lora.md", "Capitol 65 — El node LoRa al camp"),
    ("66-telegram.md", "Capitol 66 — Bot de Telegram: alertes al mòbil"),
    ("67-prometheus-alertes.md", "Capitol 67 — Prometheus i alertes avançades"),
    ("68-runbooks.md", "Capitol 68 — Runbooks: quan falla alguna cosa"),
    ("69-drp-test.md", "Capitol 69 — DRP: el dia que es crema tot"),
]

# Mòdul actiu per defecte. Es pot canviar amb --module {1|2|both}
DEFAULT_MODULE = "both"


# ---------------------------------------------------------------------------
# Parser Markdown -> llista de blocs
# ---------------------------------------------------------------------------


def split_blocks(text):
    """
    Parteix el text en blocs lògics.
    Tipus possibles: 'h1','h2','h3','h4','p','code','quote','ul','ol','hr','mermaid','table'.
    Retorna llista de dicts: {type, content (str|list)}.
    """
    lines = text.splitlines()
    blocks = []
    i = 0
    n = len(lines)

    def is_hr(line):
        s = line.strip()
        return s in ("---", "***", "___") or re.fullmatch(r"-{3,}", s) is not None

    def is_table_header(line):
        # Una taula comença amb | ... | i la línia següent és | --- | --- |
        if "|" not in line:
            return False
        return False  # es decideix en processar

    while i < n:
        line = lines[i]

        # Línia buida
        if not line.strip():
            i += 1
            continue

        # Regla horitzontal
        if is_hr(line):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Capçaleres
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            content = m.group(2).strip()
            blocks.append({"type": f"h{level}", "content": content})
            i += 1
            continue

        # Bloc de codi ```lang
        if line.lstrip().startswith("```"):
            lang = line.strip().lstrip("`").strip()
            i += 1
            buf = []
            while i < n and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            # saltar la línia de tancament
            if i < n:
                i += 1
            blocks.append({"type": "code", "lang": lang, "content": "\n".join(buf)})
            continue

        # Mermaid (molt específic: bloc de codi mermaid el tractem com a
        # text en una caixa — dibuixar un Mermaid real requereix headless
        # browser, massa per a aquesta versió).
        # Si ja l'hem detectat com a code amb lang=mermaid, el re-tractem.

        # Cita
        if line.lstrip().startswith(">"):
            buf = []
            while i < n and lines[i].lstrip().startswith(">"):
                buf.append(lines[i].lstrip()[1:].lstrip())
                i += 1
            blocks.append({"type": "quote", "content": "\n".join(buf)})
            continue

        # Llistes no ordenades
        if re.match(r"^\s*[-*+]\s+", line):
            items = []
            while i < n:
                m2 = re.match(r"^(\s*)[-*+]\s+(.*)$", lines[i])
                if not m2:
                    # línia de continuació (sagnat)
                    if lines[i].startswith("  ") or lines[i].startswith("\t"):
                        if items:
                            items[-1] += "\n" + lines[i].strip()
                            i += 1
                            continue
                    break
                indent, content = m2.group(1), m2.group(2)
                items.append(content.strip())
                i += 1
            blocks.append({"type": "ul", "items": items})
            continue

        # Llistes ordenades
        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < n:
                m2 = re.match(r"^\s*\d+\.\s+(.*)$", lines[i])
                if not m2:
                    if lines[i].startswith("  ") or lines[i].startswith("\t"):
                        if items:
                            items[-1] += "\n" + lines[i].strip()
                            i += 1
                            continue
                    break
                items.append(m2.group(1).strip())
                i += 1
            blocks.append({"type": "ol", "items": items})
            continue

        # Taules
        if "|" in line and i + 1 < n and re.match(r"^\s*\|?\s*:?-{2,}", lines[i + 1].replace("|", " | ")) or (
            "|" in line and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1])
        ):
            # confirmar separador
            sep = lines[i + 1].strip()
            if sep.startswith("|") or sep.endswith("|"):
                if re.match(r"^\|?[\s:|-]+\|?$", sep) and "-" in sep:
                    # Parsejar capçalera
                    header = [c.strip() for c in line.strip().strip("|").split("|")]
                    i += 2  # saltar capçalera i separador
                    rows = []
                    while i < n and "|" in lines[i] and lines[i].strip():
                        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                        rows.append(row)
                        i += 1
                    blocks.append({"type": "table", "header": header, "rows": rows})
                    continue

        # Paràgraf (reunir línies fins a línia buida o nova construcció)
        buf = []
        while i < n and lines[i].strip() and not re.match(r"^(#{1,6}\s+|```|>\s|[-*+]\s|\d+\.\s)", lines[i]) and not is_hr(lines[i]):
            buf.append(lines[i])
            i += 1
        if buf:
            blocks.append({"type": "p", "content": " ".join(buf)})

    return blocks


# ---------------------------------------------------------------------------
# Inline: negreta, cursiva, codi, links
# ---------------------------------------------------------------------------


def inline_md_to_html(text):
    """
    Converteix un fragment de text amb sintaxi Markdown inline a HTML
    compatible amb reportlab. Gestiona: **negreta**, *cursiva*, `codi`,
    [text](url). Escapa la resta (& < >).
    """
    # Escapar HTML
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    # Codi inline `...` (primer perquè no volem que s'interpreti res dins)
    parts = re.split(r"(`[^`\n]+`)", text)
    for idx, part in enumerate(parts):
        if part.startswith("`") and part.endswith("`"):
            parts[idx] = f'<font name="Courier" color="#1f3a5f">{part[1:-1]}</font>'
        else:
            # Negreta **...**
            part = re.sub(
                r"\*\*([^\*\n]+?)\*\*",
                r'<b>\1</b>',
                part,
            )
            # Cursiva *...*
            part = re.sub(
                r"(?<!\*)\*([^\*\n]+?)\*(?!\*)",
                r'<i>\1</i>',
                part,
            )
            # Enllaços [text](url)
            part = re.sub(
                r"\[([^\]]+)\]\(([^)]+)\)",
                r'<font color="#1f4e79"><u>\1</u></font>',
                part,
            )
            parts[idx] = part
    return "".join(parts)


# ---------------------------------------------------------------------------
# Generació PDF amb reportlab
# ---------------------------------------------------------------------------


def build_pdf(chapters, out_pdf_path, module_label, module_subtitle):
    print(f"[pdf] Construint {out_pdf_path.name} ({module_label}) ...")

    # Estils
    styles = getSampleStyleSheet()

    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=colors.HexColor("#1a1a1a"),
        spaceAfter=6,
        spaceBefore=0,
        alignment=0,  # esquerra
    )

    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        textColor=colors.HexColor("#1f3a5f"),
        spaceBefore=18,
        spaceAfter=12,
        keepWithNext=True,
    )

    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#1f3a5f"),
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True,
    )

    h3 = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor("#2c5d8a"),
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=True,
    )

    h4 = ParagraphStyle(
        "H4",
        parent=styles["Heading4"],
        fontName="Helvetica-BoldOblique",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#3a3a3a"),
        spaceBefore=8,
        spaceAfter=4,
        keepWithNext=True,
    )

    code_style = ParagraphStyle(
        "Code",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#1a1a1a"),
        backColor=colors.HexColor("#f4f4f0"),
        borderColor=colors.HexColor("#cccccc"),
        borderWidth=0.5,
        borderPadding=6,
        leftIndent=0,
        rightIndent=0,
        spaceBefore=4,
        spaceAfter=8,
    )

    quote_style = ParagraphStyle(
        "Quote",
        parent=body,
        leftIndent=14,
        rightIndent=8,
        textColor=colors.HexColor("#3a3a3a"),
        fontName="Helvetica-Oblique",
        borderColor=colors.HexColor("#1f3a5f"),
        borderWidth=0,
        borderPadding=0,
        spaceBefore=4,
        spaceAfter=8,
    )

    cover_title = ParagraphStyle(
        "CoverTitle",
        fontName="Helvetica-Bold",
        fontSize=36,
        leading=42,
        textColor=colors.HexColor("#1f3a5f"),
        alignment=1,
        spaceAfter=20,
    )

    cover_sub = ParagraphStyle(
        "CoverSub",
        fontName="Helvetica-Oblique",
        fontSize=14,
        leading=20,
        textColor=colors.HexColor("#3a3a3a"),
        alignment=1,
        spaceAfter=14,
    )

    cover_meta = ParagraphStyle(
        "CoverMeta",
        fontName="Helvetica",
        fontSize=11,
        leading=16,
        textColor=colors.HexColor("#555555"),
        alignment=1,
        spaceAfter=4,
    )

    # Doc amb peu de pàgina + capçalera
    doc = BaseDocTemplate(
        str(out_pdf_path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2.2 * cm,
        bottomMargin=2.2 * cm,
        title=f"BernatLab — Manual tècnic pràctic ({module_label})",
        author="Bernat",
        subject=f"BernatLab — {module_subtitle}",
    )

    def on_page(canvas, doc_):
        canvas.saveState()
        # Capçalera
        canvas.setFont("Helvetica", 8.5)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawString(2 * cm, A4[1] - 1.2 * cm, "BernatLab · Manual tècnic pràctic")
        canvas.drawRightString(
            A4[0] - 2 * cm, A4[1] - 1.2 * cm, module_label
        )
        # peu
        canvas.setFont("Helvetica", 8.5)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawCentredString(
            A4[0] / 2.0, 1.0 * cm, f"Pàgina {doc_.page}"
        )
        # línia sota capçalera
        canvas.setStrokeColor(colors.HexColor("#cccccc"))
        canvas.setLineWidth(0.4)
        canvas.line(2 * cm, A4[1] - 1.4 * cm, A4[0] - 2 * cm, A4[1] - 1.4 * cm)
        canvas.restoreState()

    frame = Frame(
        doc.leftMargin,
        doc.bottomMargin,
        doc.width,
        doc.height,
        id="normal",
        showBoundary=0,
    )
    template = PageTemplate(id="cover", frames=frame, onPage=on_page)
    doc.addPageTemplates([template])

    story = []

    # --- Portada ---
    story.append(Spacer(1, 4.5 * cm))
    story.append(Paragraph("BernatLab", cover_title))
    story.append(Paragraph(module_subtitle.split(" · ")[0] if " · " in module_subtitle else module_subtitle, cover_sub))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Raspberry Pi · Docker · IA · LoRa · Hort Osona", cover_sub))
    story.append(Spacer(1, 5 * cm))
    story.append(Paragraph(module_label, cover_meta))
    story.append(Paragraph("Bernat — 8 de juliol del 2026", cover_meta))
    story.append(Spacer(1, 0.6 * cm))
    story.append(Paragraph("Document tècnic de consulta personal", cover_meta))
    story.append(PageBreak())

    # --- Pròleg ---
    if module_label == "Mòdul 1":
        proleg_text = (
            "Aquest és el primer mòdul del manual tècnic del BernatLab. Cobreix els fonaments: "
            "comprendre la Raspberry Pi, administrar Linux, configurar xarxa i SSH, desplegar "
            "contenidors amb Docker, gestionar-los amb Portainer, monitorar-los amb Uptime Kuma, "
            "presentar-los amb Homepage, versionar-ho tot amb Git, i dibuixar la full de ruta "
            "dels pròxims mesos."
        )
        idx_lines = [
            "1. Què és BernatLab",
            "2. La Raspberry Pi 4 per dins",
            "3. Linux per administrar un servidor",
            "4. Xarxa, SSH i Tailscale",
            "5. Docker des de zero",
            "6. Portainer",
            "7. Uptime Kuma",
            "8. Homepage",
            "9. Git i documentació",
            "10. Full de ruta del BernatLab",
        ]
    else:
        proleg_text = (
            "Aquest és el segon mòdul del manual tècnic del BernatLab. Cobreix tota la cadena "
            "de sensors i visualització: el protocol MQTT, el broker Mosquitto, l'esquema de "
            "publicació dels sensors, la base de dades InfluxDB, l'agent Telegraf, la programació "
            "visual amb Node-RED, la visualització amb Grafana, l'API REST amb FastAPI, "
            "la integració amb la web pública Hort Osona i l'operativa del dia a dia."
        )
        idx_lines = [
            "11. Del Mòdul 1 al M2: què construïm",
            "12. MQTT des de zero",
            "13. Mosquitto al BernatLab",
            "14. Publicar dades: els sensors",
            "15. InfluxDB: base de dades de sèries temporals",
            "16. Telegraf: el pont",
            "17. Node-RED: programació visual",
            "18. Fluxos pràctics",
            "19. Grafana: visualitzar les dades",
            "20. API pública: servir les dades al món",
            "21. Integració amb Hort Osona",
            "22. Operativa: còpies, alertes, escalat",
        ]

    story.append(Paragraph("Sobre aquest manual", h1))
    story.append(Paragraph(proleg_text, body))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Com es llegeix", h2))
    story.append(Paragraph(
        "En ordre, si parteixes de zero. Per capítols, si ja tens una base i vols aprofundir. "
        "Cada capítol segueix la mateixa estructura: teoria, aplicació al BernatLab, esquemes, "
        "comandes útils, errors habituals, exercicis pràctics i resum final.",
        body,
    ))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Índex de capítols", h2))
    for line in idx_lines:
        story.append(Paragraph(f"• {line}", body))
    story.append(PageBreak())

    # --- Capítols ---
    for filename, title in chapters:
        path = CAP_DIR / filename
        if not path.exists():
            print(f"  AVÍS: no trobo {path}")
            continue
        text = path.read_text(encoding="utf-8")
        # La primera línia del capítol sol ser "# Títol" — la ignorem perquè
        # ja posem el nostre propi títol a la capçalera del capítol.
        if text.lstrip().startswith("# "):
            # saltar fins al primer paràgraf o capçalera de nivell >1
            lines = text.splitlines()
            for k, ln in enumerate(lines):
                if ln.startswith("# ") and ln.strip().lstrip("# ").strip():
                    text = "\n".join(lines[k + 1 :])
                    break

        story.append(Paragraph(title, h1))
        story.append(Spacer(1, 0.2 * cm))

        blocks = split_blocks(text)
        for b in blocks:
            t = b["type"]
            if t == "h1":
                story.append(Paragraph(inline_md_to_html(b["content"]), h2))
            elif t == "h2":
                story.append(Paragraph(inline_md_to_html(b["content"]), h2))
            elif t == "h3":
                story.append(Paragraph(inline_md_to_html(b["content"]), h3))
            elif t == "h4":
                story.append(Paragraph(inline_md_to_html(b["content"]), h4))
            elif t == "p":
                story.append(Paragraph(inline_md_to_html(b["content"]), body))
            elif t == "quote":
                # Cada línia de la cita va en un paràgraf amb estil cita
                for ql in b["content"].splitlines():
                    if ql.strip():
                        story.append(Paragraph(ql, quote_style))
                    else:
                        story.append(Spacer(1, 4))
            elif t == "code":
                lang = b.get("lang", "")
                if lang == "mermaid":
                    # Bloc Mermaid: el mostrem com a text en una caixa amb
                    # una nota indicant que és un esquema.
                    code_text = b["content"]
                    # Carregar estil amb mida més petita si és llarg
                    style = ParagraphStyle(
                        "Mermaid",
                        parent=code_style,
                        fontSize=7.8,
                        leading=9.5,
                    )
                    note = Paragraph(
                        '<font color="#7a3a00"><i>Esquema Mermaid (text):</i></font>',
                        body,
                    )
                    story.append(note)
                    # Preformatted amb salts de línia preservats
                    pre = Preformatted(code_text, style)
                    story.append(pre)
                else:
                    code_text = b["content"]
                    pre = Preformatted(code_text, code_style)
                    story.append(pre)
            elif t == "ul":
                for it in b["items"]:
                    story.append(Paragraph("• " + inline_md_to_html(it), body))
                story.append(Spacer(1, 2))
            elif t == "ol":
                for n, it in enumerate(b["items"], 1):
                    story.append(Paragraph(f"{n}. " + inline_md_to_html(it), body))
                story.append(Spacer(1, 2))
            elif t == "hr":
                # regla visual: línia + espai
                story.append(Spacer(1, 0.2 * cm))
                # usem un Paragraph amb un border bottom
                rule = Paragraph(
                    '<hr width="100%" size="1" color="#cccccc"/>',
                    ParagraphStyle("Rule", fontSize=2, leading=2),
                )
                story.append(rule)
                story.append(Spacer(1, 0.2 * cm))
            elif t == "table":
                header = b["header"]
                rows = b["rows"]
                data = [header] + rows
                # calcular amplades
                ncols = len(header)
                avail = doc.width
                col_w = [avail / ncols] * ncols
                t = Table(data, colWidths=col_w, repeatRows=1)
                t.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f3a5f")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE", (0, 0), (-1, -1), 9),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                            ("TOPPADDING", (0, 0), (-1, -1), 5),
                            ("LEFTPADDING", (0, 0), (-1, -1), 5),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#888888")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#ffffff"), colors.HexColor("#f4f4f0")]),
                        ]
                    )
                )
                story.append(t)
                story.append(Spacer(1, 0.3 * cm))
        # forçar salt de pàgina entre capítols (excepte l'últim)
        if filename != chapters[-1][0]:
            story.append(PageBreak())

    doc.build(story)
    size = out_pdf_path.stat().st_size
    print(f"[pdf] OK: {out_pdf_path} ({size:,} bytes)")


# ---------------------------------------------------------------------------
# Generació DOCX amb python-docx
# ---------------------------------------------------------------------------


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_code_block(doc, text, lang=""):
    # quadre gris al voltant
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F0")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "CCCCCC")
        pBdr.append(b)
    pPr.append(pBdr)
    for i, line in enumerate(text.splitlines()):
        if i > 0:
            p.add_run("\n")
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_quote(doc, text):
    for line in text.splitlines():
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(line)
        run.italic = True
        run.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)


def add_inline_runs(paragraph, text):
    """
    Aplica negreta, cursiva, codi inline i links dins d'un paràgraf DOCX.
    Fa servir un escàner d'estat perquè el text pot contenir **, *, `, [.
    """
    # Patrons: `codi`, **negreta**, *cursiva*, [text](url)
    pattern = re.compile(
        r"(`[^`\n]+`|\*\*[^\*\n]+?\*\*|\*[^\*\n]+?\*|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
        chunk = m.group(0)
        if chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("["):
            mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", chunk)
            if mm:
                run = paragraph.add_run(mm.group(1))
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                run.underline = True
            else:
                run = paragraph.add_run(chunk)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx(chapters, out_docx_path, module_label, module_subtitle):
    print(f"[docx] Construint {out_docx_path.name} ({module_label}) ...")

    doc = Document()

    # Estils base
    for style_name, size, bold, color in [
        ("Normal", 11, False, None),
    ]:
        s = doc.styles[style_name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        if color:
            s.font.color.rgb = color

    # Capçalera i peu de pàgina
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "BernatLab · Manual tècnic pràctic"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = module_label
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in fp.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # --- Portada ---
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BernatLab")
    r.font.size = Pt(40)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(module_subtitle.split(" · ")[0] if " · " in module_subtitle else module_subtitle)
    r.font.size = Pt(16)
    r.italic = True
    r.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Raspberry Pi · Docker · IA · LoRa · Hort Osona")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(8):
        doc.add_paragraph()

    for line in [
        module_label,
        "Bernat — 8 de juliol del 2026",
        "Document tècnic de consulta personal",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # Pròleg + índex
    h = doc.add_heading("Sobre aquest manual", level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    if module_label == "Mòdul 1":
        proleg_text = (
            "Aquest és el primer mòdul del manual tècnic del BernatLab. Cobreix els fonaments: "
            "comprendre la Raspberry Pi, administrar Linux, configurar xarxa i SSH, desplegar "
            "contenidors amb Docker, gestionar-los amb Portainer, monitorar-los amb Uptime Kuma, "
            "presentar-los amb Homepage, versionar-ho tot amb Git, i dibuixar la full de ruta "
            "dels pròxims mesos."
        )
        idx_lines = [
            "1. Què és BernatLab",
            "2. La Raspberry Pi 4 per dins",
            "3. Linux per administrar un servidor",
            "4. Xarxa, SSH i Tailscale",
            "5. Docker des de zero",
            "6. Portainer",
            "7. Uptime Kuma",
            "8. Homepage",
            "9. Git i documentació",
            "10. Full de ruta del BernatLab",
        ]
    else:
        proleg_text = (
            "Aquest és el segon mòdul del manual tècnic del BernatLab. Cobreix tota la cadena "
            "de sensors i visualització: el protocol MQTT, el broker Mosquitto, l'esquema de "
            "publicació dels sensors, la base de dades InfluxDB, l'agent Telegraf, la programació "
            "visual amb Node-RED, la visualització amb Grafana, l'API REST amb FastAPI, "
            "la integració amb la web pública Hort Osona i l'operativa del dia a dia."
        )
        idx_lines = [
            "11. Del Mòdul 1 al M2: què construïm",
            "12. MQTT des de zero",
            "13. Mosquitto al BernatLab",
            "14. Publicar dades: els sensors",
            "15. InfluxDB: base de dades de sèries temporals",
            "16. Telegraf: el pont",
            "17. Node-RED: programació visual",
            "18. Fluxos pràctics",
            "19. Grafana: visualitzar les dades",
            "20. API pública: servir les dades al món",
            "21. Integració amb Hort Osona",
            "22. Operativa: còpies, alertes, escalat",
        ]
    p = doc.add_paragraph(proleg_text)

    doc.add_heading("Com es llegeix", level=2)
    doc.add_paragraph(
        "En ordre, si parteixes de zero. Per capítols, si ja tens una base i vols aprofundir. "
        "Cada capítol segueix la mateixa estructura: teoria, aplicació al BernatLab, esquemes, "
        "comandes útils, errors habituals, exercicis pràctics i resum final."
    )

    doc.add_heading("Índex de capítols", level=2)
    for line in idx_lines:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_page_break()

    # Capítols
    for filename, title in chapters:
        path = CAP_DIR / filename
        if not path.exists():
            print(f"  AVÍS: no trobo {path}")
            continue
        text = path.read_text(encoding="utf-8")
        if text.lstrip().startswith("# "):
            lines = text.splitlines()
            for k, ln in enumerate(lines):
                if ln.startswith("# ") and ln.strip().lstrip("# ").strip():
                    text = "\n".join(lines[k + 1 :])
                    break

        h = doc.add_heading(title, level=1)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

        blocks = split_blocks(text)
        for b in blocks:
            t = b["type"]
            if t == "h1":
                h = doc.add_heading(b["content"], level=2)
                for r in h.runs:
                    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            elif t == "h2":
                h = doc.add_heading(b["content"], level=2)
                for r in h.runs:
                    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            elif t == "h3":
                h = doc.add_heading(b["content"], level=3)
                for r in h.runs:
                    r.font.color.rgb = RGBColor(0x2C, 0x5D, 0x8A)
            elif t == "h4":
                h = doc.add_heading(b["content"], level=4)
            elif t == "p":
                p = doc.add_paragraph()
                add_inline_runs(p, b["content"])
            elif t == "quote":
                add_quote(doc, b["content"])
            elif t == "code":
                lang = b.get("lang", "")
                if lang == "mermaid":
                    p = doc.add_paragraph()
                    r = p.add_run("Esquema Mermaid (text):")
                    r.italic = True
                    r.font.color.rgb = RGBColor(0x7A, 0x3A, 0x00)
                    add_code_block(doc, b["content"], lang)
                else:
                    add_code_block(doc, b["content"], lang)
            elif t == "ul":
                for it in b["items"]:
                    p = doc.add_paragraph(style="List Bullet")
                    add_inline_runs(p, it)
            elif t == "ol":
                for it in b["items"]:
                    p = doc.add_paragraph(style="List Number")
                    add_inline_runs(p, it)
            elif t == "hr":
                add_horizontal_rule(doc)
            elif t == "table":
                header = b["header"]
                rows = b["rows"]
                table = doc.add_table(rows=1 + len(rows), cols=len(header))
                table.style = "Light Grid Accent 1"
                hdr = table.rows[0].cells
                for i, c in enumerate(header):
                    hdr[i].text = c
                    for p in hdr[i].paragraphs:
                        for r in p.runs:
                            r.bold = True
                for ri, row in enumerate(rows, 1):
                    for ci, val in enumerate(row):
                        if ci < len(table.rows[ri].cells):
                            cell = table.rows[ri].cells[ci]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            add_inline_runs(p, val)
        if filename != chapters[-1][0]:
            doc.add_page_break()

    doc.save(str(out_docx_path))
    size = out_docx_path.stat().st_size
    print(f"[docx] OK: {out_docx_path} ({size:,} bytes)")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main():
    if not CAP_DIR.exists():
        print(f"ERROR: no trobo la carpeta {CAP_DIR}")
        sys.exit(1)

    # Argument --module {1|2|3|both|all}
    module = DEFAULT_MODULE
    if len(sys.argv) > 1:
        module = sys.argv[1]
        if module not in ("1", "2", "3", "4", "5", "6", "7", "both", "all"):
                print(f"ERROR: --module ha de ser 1, 2, 3, 4, both o all (rebut: {module})")
                sys.exit(1)

    if module in ("1", "both", "all"):
        out_pdf = ROOT / "output" / "BernatLab_Manual_Modul_1.pdf"
        out_docx = ROOT / "output" / "BernatLab_Manual_Modul_1.docx"
        build_pdf(CHAPTERS_M1, out_pdf, "Mòdul 1", "Fonaments, contenidors i pràctica")
        build_docx(CHAPTERS_M1, out_docx, "Mòdul 1", "Fonaments, contenidors i pràctica")

    if module in ("2", "both", "all"):
        out_pdf = ROOT / "output" / "BernatLab_Manual_Modul_2.pdf"
        out_docx = ROOT / "output" / "BernatLab_Manual_Modul_2.docx"
        build_pdf(CHAPTERS_M2, out_pdf, "Mòdul 2", "Sensors, dades i visualització")
        build_docx(CHAPTERS_M2, out_docx, "Mòdul 2", "Sensors, dades i visualització")

    if module in ("3", "all"):
        out_pdf = ROOT / "output" / "BernatLab_Manual_Modul_3.pdf"
        out_docx = ROOT / "output" / "BernatLab_Manual_Modul_3.docx"
        build_pdf(CHAPTERS_M3, out_pdf, "Mòdul 3", "LoRa, sensors remots i xarxa de camp")
        build_docx(CHAPTERS_M3, out_docx, "Mòdul 3", "LoRa, sensors remots i xarxa de camp")

    if module in ("4", "all"):
        out_pdf = ROOT / "output" / "BernatLab_Manual_Modul_4.pdf"
        out_docx = ROOT / "output" / "BernatLab_Manual_Modul_4.docx"
        build_pdf(CHAPTERS_M4, out_pdf, "Mòdul 4", "IA local amb Ollama: RAG, veu i privadesa")
        build_docx(CHAPTERS_M4, out_docx, "Mòdul 4", "IA local amb Ollama: RAG, veu i privadesa")

    if module in ("5", "all"):
        out_pdf = ROOT / "output" / "BernatLab_Manual_Modul_5.pdf"
        out_docx = ROOT / "output" / "BernatLab_Manual_Modul_5.docx"
        build_pdf(CHAPTERS_M5, out_pdf, "Mòdul 5", "Seguretat i còpies de seguretat")
        build_docx(CHAPTERS_M5, out_docx, "Mòdul 5", "Seguretat i còpies de seguretat")

    if module in ("6", "all"):
        out_pdf = ROOT / "output" / "BernatLab_Manual_Modul_6.pdf"
        out_docx = ROOT / "output" / "BernatLab_Manual_Modul_6.docx"
        build_pdf(CHAPTERS_M6, out_pdf, "Mòdul 6", "Operativa 24/7, monitoratge i manteniment")
        build_docx(CHAPTERS_M6, out_docx, "Mòdul 6", "Operativa 24/7, monitoratge i manteniment")

    if module in ("7", "all"):
        out_pdf = ROOT / "output" / "BernatLab_Manual_Modul_7.pdf"
        out_docx = ROOT / "output" / "BernatLab_Manual_Modul_7.docx"
        build_pdf(CHAPTERS_M7, out_pdf, "Mòdul 7", "Hort Osona en acció: implementar el BernatLab pas a pas")
        build_docx(CHAPTERS_M7, out_docx, "Mòdul 7", "Hort Osona en acció: implementar el BernatLab pas a pas")

    print("[fet] Generació completada.")


if __name__ == "__main__":
    main()
